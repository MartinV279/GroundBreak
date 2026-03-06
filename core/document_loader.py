from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    path: Path
    name: str
    text: str


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def discover_documents(root: Path) -> List[Path]:
    paths: List[Path] = []
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS:
            paths.append(p)
    return paths


def _load_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def _load_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts: List[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts)
    except Exception:
        return ""


def _load_docx(path: Path) -> str:
    try:
        doc = DocxDocument(str(path))
        paras = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(paras)
    except Exception:
        return ""


def load_documents(root: Path) -> List[LoadedDocument]:
    docs: List[LoadedDocument] = []
    for p in discover_documents(root):
        ext = p.suffix.lower()
        if ext == ".txt":
            text = _load_txt(p)
        elif ext == ".pdf":
            text = _load_pdf(p)
        elif ext == ".docx":
            text = _load_docx(p)
        else:
            continue
        if not text.strip():
            continue
        docs.append(LoadedDocument(path=p, name=p.name, text=text))
    return docs

